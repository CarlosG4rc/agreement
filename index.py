from tkinter import *
from tkinter import filedialog as FileDialog
from io import open
import openpyxl
import docx
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def crear():
   ss = openpyxl.load_workbook(ex.get() + ".xlsx")
   sheet = ss.get_sheet_by_name(hoja.get())
   
   for i in sheet.iter_rows(max_row=0):
      n = len(i)
   #Extraemos datos de contrato desde Excel
   for j in range(6, n):
      nombre = sheet.cell(row = j, column = 1).value
      nacionalidad = sheet.cell(row = j, column = 2).value
      domicilio = sheet.cell(row = j, column = 3).value
      curp = sheet.cell(row = j, column = 4).value
      rfc = sheet.cell(row = j, column = 5).value
      start = sheet.cell(row = j, column = 6).value
      final = sheet.cell(row = j, column = 7).value
      hrs = sheet.cell(row = j, column = 8).value
      hrs = str(hrs)
      importe = sheet.cell(row = j, column = 9).value
      importe = str(importe)
      importe_letra = sheet.cell(row = j, column = 10).value
      antiguedad = sheet.cell(row = j, column = 11).value

      #Creamos el docx
      doc = docx.Document()
      paragraph = doc.add_paragraph()
      run = paragraph.add_run("CONTRATO INDIVIDUAL DE TRABAJO POR TIEMPO DETERMINADO PARA PROFESORES")
      font = run.font
      paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
      font.color.rgb = RGBColor(0x00, 0x00, 0x00)
      font.name = 'Tahoma'
      font.size = Pt(11)
      font.bold = True

      paragraph = doc.add_paragraph()
      run = paragraph.add_run("CONTRATO INDIVIDUAL DE TRABAJO POR TIEMPO DETERMINADO PARA MAESTROS QUE CELEBRAN POR UNA PARTE EL COLEGIO ")
      font = run.font
      paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)
      
      run1 = paragraph.add_run('"INSTITUTO FRANCISCO POSSENTI, A. C." ')
      run1.bold = TRUE
      font = run1.font
      paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run = paragraph.add_run("CON DOMICILIO EN AV. TOLUCA No. 621 COL. OLIVAR DE LOS PADRES DEL. ALVARO OBREGON C. P. 01780 REPRESENTADO POR EL C. ")
      font = run.font
      paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run1 = paragraph.add_run("J. ANTONIO BARRIENTOS RODRIGUEZ ")
      run1.bold = TRUE
      font = run1.font
      paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run = paragraph.add_run("A QUIEN EN LO SUCESIVO SE DENOMINARA EL PATRON, Y POR LA OTRA. " + nombre + " DE NACIONALIDAD "+ nacionalidad + " CON DOMICILIO "+ domicilio + ", A QUIEN EN ADELANTE SE DENOMINARA EL TRABAJADOR, DE ACUERDO CON LAS SIGUIENTES:")
      font = run.font
      paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      clau = doc.add_paragraph()
      run2 = clau.add_run("C L A U S U L A S ")
      font = run2.font
      clau.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
      font.color.rgb = RGBColor(0x00, 0x00, 0x00)
      font.name = 'Arial'
      font.size = Pt(11)
      font.bold = True

      paragraph2 = doc.add_paragraph()
      run3 = paragraph2.add_run("PRIMERA.- ")
      run3.bold = TRUE
      font = run3.font
      paragraph2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run3 = paragraph2.add_run("El (a) Profesor (a) manifiesta, bajo protesta de decir verdad, que tiene la Clave Única de Registro de Población " + curp + " y el Registro Federal de Contribuyentes "+ rfc + " que tiene  la capacidad, aptitudes, facultades y conocimientos necesarios para desempeñar el trabajo que se le encomienda, así como  la documentación completa y actualizada por la Secretaria de Educación Publica y/o la UNAM, así como  a las disposiciones señaladas por los artículos 42 fracción VII, de la nueva Ley Federal  del Trabajo publicada en el Diario Oficial de la Federación el día 30 de noviembre  del 2012  que se requiere así como está de acuerdo en que el no cumplir con cualquiera de estos requisitos será causa suficiente para que el patrón le rescinda su contrato de trabajo en el momento que tenga conocimiento de la carencia de alguna de esta condiciones, así mismo se compromete a que en caso de que el profesor (a)  cambie de domicilio durante la vigencia del presente contrato notificará por escrito al patrón dentro de los  cinco días siguientes que cambie de domicilio. ")
      font = run3.font
      paragraph2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraph3 = doc.add_paragraph()
      run3 = paragraph3.add_run("SEGUNDA.- ")
      run3.bold = TRUE
      font = run3.font
      paragraph3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run3 = paragraph3.add_run("Este contrato por exigencias expresas de la Secretaría de Educación Pública se celebra por tiempo determinado, el cual se precisa en el  Acuerdo ")
      font = run3.font
      paragraph3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)
      
      run3 = paragraph3.add_run("14/072020 ")
      run3.bold = TRUE
      font = run3.font
      paragraph2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run3 = paragraph3.add_run("de la Secretaría de Educación Pública publicado en el Diario Oficial del ")
      font = run3.font
      paragraph3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run3 = paragraph3.add_run("03 DE AGOSTO DEL 2020, ")
      run3.bold = TRUE
      font = run3.font
      paragraph2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run3 = paragraph3.add_run("y sólo podrá modificarse, rescindirse o terminarse en los casos y condiciones especificados en la Ley Federal del Trabajo, o por aquellas autoridades que en su momento cuenten con facultades suficientes para modificar, rescindir o dar por terminado el presente contrato. ")
      font = run3.font
      paragraph3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraph4 = doc.add_paragraph()
      run3 = paragraph4.add_run("TERCERA.- ")
      run3.bold = TRUE
      font = run3.font
      paragraph4.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run3 = paragraph4.add_run("El Patrón y el (a) Profesor (a), convienen expresamente y con fundamento en el Art. 47 Fracción I de la Ley Federal del Trabajo, que dentro de los primeros treinta días o cuando el patrón tenga conocimiento de la carencia o incumplimiento de alguna de las condiciones básicas requeridas para desempeñar el trabajo contratado, se podrá rescindir este Contrato de Trabajo sin responsabilidad para el Patrón. ")
      font = run3.font
      paragraph4.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraph5 = doc.add_paragraph()
      run3 = paragraph5.add_run("CUARTA.- ")
      run3.bold = TRUE
      font = run3.font
      paragraph5.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run3 = paragraph5.add_run("El (a) Profesor (a) se obliga a prestar sus servicios personales al INSTITUTO, bajo su dirección, dependencia y subordinación, las cuales consistirán precisamente en: ")
      font = run3.font
      paragraph5.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraph6 = doc.add_paragraph()
      run3 = paragraph6.add_run("Proporcionar personalmente, a los alumnos que se le indiquen o le sean asignados enseñanza eficiente durante el tiempo determinado del ciclo Escolar vigente, y como lo dispone el artículo 56 Bis de la nueva Ley Federal del Trabajo publicada en el Diario Oficial de la Federación el día 30 de noviembre del 2012,  sujetándose a los programas y planes de estudio correspondientes que le sean entregados por el INSTITUTO, debidamente autorizados. ")
      font = run3.font
      paragraph6.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraph6 = doc.add_paragraph()
      run3 = paragraph6.add_run("El presente contrato se celebra por un tiempo que será de " + start + " y termina " + final)
      font = run3.font
      paragraph6.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraph7 = doc.add_paragraph()
      run3 = paragraph7.add_run("QUINTA.- ")
      run3.bold = TRUE
      font = run3.font
      paragraph7.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run3 = paragraph7.add_run("Los servicios contratados se estipulan en forma enunciativa y no limitativa; por tanto, el (a) Profesor (a)  se obliga a desempeñar  todas las labores anexas o conexas con su obligación principal y las demás que le ordene el Patrón o sus representantes, tales como guardias escolares, cursos de verano, de capacitación de reprogramación de estudios, exámenes extraordinarios, exámenes de diagnóstico, noche colonial, exámenes de admisión, ofrenda de día de muertos, posada navideña etc, cuya retribución económica está convenida y comprendida en la Cláusula Décima Primera del presente contrato.")
      font = run3.font
      paragraph7.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraph8 = doc.add_paragraph()
      run3 = paragraph8.add_run("De la misma manera y solo para el caso de que sea aplicable y derivado de  un caso  fortuito o fuerza mayor las modificaciones al presente contrato referidas en las reformas al artículo  311 de la Ley Federal del Trabajo capitulo XII Bis en materia de Teletrabajo publicada en el diario Oficial de la federación el día 11 de enero del 2021.  ")
      font = run3.font
      paragraph8.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraph9 = doc.add_paragraph()
      run3 = paragraph9.add_run("La desobediencia a las órdenes o indicaciones del Patrón o sus representantes para el cumplimiento del trabajo contratado, será causa de rescisión sin responsabilidad para el Patrón.")
      font = run3.font
      paragraph9.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraphA = doc.add_paragraph()
      run3 = paragraphA.add_run("SEXTA.- ")
      run3.bold = TRUE
      font = run3.font
      paragraphA.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run3 = paragraphA.add_run("Los servicios objeto de la relación de trabajo deben prestarse en el lugar o en los lugares que designe el Patrón o sus representantes, quedando convenido que éste tendrá derecho de cambiar el lugar de trabajo del (a) profesor (a) cuando se estime pertinente o necesario, siempre y cuando dicho cambio no se traduzca en una merma de su remuneración para el (la) mismo (a),esto incluye para aquellos casos donde las autoridades competentes establezcan el cierre de la fuente de trabajo derivado de caso fortuito o fuerza mayor así como las posibles modificaciones al presente contrato según las reformas al artículo 311 de la ley Federal del Trabajo capitulo XII Bis en materia de teletrabajo publicada en el diario oficial de la federación el día 11 de enero del 2021. ")
      font = run3.font
      paragraphA.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraphB = doc.add_paragraph()
      run3 = paragraphB.add_run("SEPTIMA.- ")
      run3.bold = TRUE
      font = run3.font
      paragraphB.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run3 = paragraphB.add_run("La duración de la jornada de trabajo será de: ")
      font = run3.font
      paragraphA.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run3 = paragraphB.add_run(hrs)
      run3.bold = TRUE
      font = run3.font
      paragraphB.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run3 = paragraphB.add_run(" horas, ")
      run3.bold = TRUE
      font = run3.font
      paragraphB.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run3 = paragraphB.add_run("según horario anexo. El (la) Profesor (a) está de acuerdo en que deberá asistir los días que sean necesarios para las distintas actividades que se precisan en la cláusula quinta del presente contrato. El pago correspondiente a esta jornada de trabajo, está ya integrado en el sueldo convenido que se indica  en la cláusula Décima del presente instrumento.")
      font = run3.font
      paragraphB.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraphC = doc.add_paragraph()
      run3 = paragraphC.add_run("OCTAVA.- ")
      run3.bold = TRUE
      font = run3.font
      paragraphC.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run3 = paragraphC.add_run("El (a) Profesor (a) se obliga a desempeñar sus labores con la intensidad, cuidado y esmero apropiados en la forma, tiempo y lugar a que se refiere este Contrato y el Reglamento Interior de Trabajo. El incumplimiento de esta disposición se considera falta de probidad del (a) Profesor (a) y, de ocurrir, se sancionará con la rescisión  del Contrato sin responsabilidad para el Patrón.")
      font = run3.font
      paragraphC.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraphD = doc.add_paragraph()
      run3 = paragraphD.add_run("NOVENA.- ")
      run3.bold = TRUE
      font = run3.font
      paragraphD.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run3 = paragraphD.add_run("El (a) Profesor (a) está obligado a checar en el reloj o firmar las listas de asistencia a la entrada y salida de sus labores, el incumplimiento de este requisito se considerará como una desobediencia para todos los efectos legales a que haya lugar,  el retiro del Profesor de su lugar de Trabajo durante su jornada de labores, sin autorización, será considerado como una desobediencia y por lo tanto abandono de trabajo.")
      font = run3.font
      paragraphD.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraphE = doc.add_paragraph()
      run3 = paragraphE.add_run("El registrar entradas o salidas por otra persona, será causa de rescisión del presente contrato sin responsabilidad para el Patrón.")
      font = run3.font
      paragraphE.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraphF = doc.add_paragraph()
      run3 = paragraphF.add_run("DECIMA.- ")
      run3.bold = TRUE
      font = run3.font
      paragraphF.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run3 = paragraphF.add_run("Se conviene como salario nominal, que el Patrón deberá pagar por los trabajos personales que reciba del(a)Profesor(a),la cantidad de: $ "+ importe + " ( " + importe_letra + " 00/100 M. N.) Dicho salario mensual será cubierto por mitad al (a) profesor (a) después de sumar las prestaciones y restar los impuestos correspondientes, cada día quince y último de cada mes mediante depósito bancario  tal y como lo dispone el artículo 101 de la nueva Ley Federal  del Trabajo publicada  en el Diario Oficial de la Federación  el día 30 de noviembre del 2012, en las oficinas de la Escuela, en efectivo o en cheque, o mediante algún otro sistema de pago que las partes estimen adecuado y seguro.") 
      font = run3.font
      paragraphF.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraphG = doc.add_paragraph()
      run3 = paragraphG.add_run("Además del salario nominal mencionado y de todas las prestaciones que establece la Ley Federal del Trabajo como mínimas EL INSTITUTO otorgará las siguientes prestaciones.") 
      font = run3.font
      paragraphG.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraphH = doc.add_paragraph()
      run3 = paragraphH.add_run("10% del salario nominal por concepto de premios por asistencia siempre y cuando el profesor no falte y  asista   a  lo convocado por el Instituto. Estos premios se entregarán en efectivo en los mismos plazos y condiciones que el salario nominal y según el Reglamento Interno de Trabajo.") 
      font = run3.font
      paragraphH.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraphI = doc.add_paragraph()
      run3 = paragraphI.add_run("10% del salario nominal por concepto de premio de puntualidad. Esta prestación la tendrá el profesor cuando llegue al Instituto 5 minutos antes de la primera hora-clase según su horario.  También se entregará en efectivo en los mismos plazos y condiciones que el salario nominal y según el Reglamento Interno de Trabajo.") 
      font = run3.font
      paragraphI.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraphJ = doc.add_paragraph()
      run3 = paragraphJ.add_run("12% del salario nominal por concepto de vales de despensa como concepto de previsión social según el Plan de Previsión  del Instituto Francisco, Possenti,A.C. los cuales se entregarán el día 28 de cada mes, en monedero electrónico.") 
      font = run3.font
      paragraphJ.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraphK = doc.add_paragraph()
      run3 = paragraphK.add_run("13% del salario nominal por concepto de aportación a un fondo de ahorro, que junto con un porcentaje igual que se retendrá al trabajador cada quincena, se depositará en una cuenta bancaria y se retirará al final del ciclo escolar") 
      font = run3.font
      paragraphK.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraphL = doc.add_paragraph()
      run3 = paragraphL.add_run("Los préstamos que se otorguen a los trabajadores, serán de acuerdo a los lineamientos que regulan el Fondo de Ahorro.") 
      font = run3.font
      paragraphL.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraphM = doc.add_paragraph()
      run3 = paragraphM.add_run("En las cantidades anteriores queda comprendido el pago de los séptimos días, los días de descanso obligatorio, las vacaciones, los décimos sextos días del mes, así como los conceptos señalados en la Cláusula Quinta del presente contrato, así como las labores conexas o complementarias que desempeña de acuerdo a su labor principal tal y como lo dispone el artículo 56 Bis de la nueva Ley Federal del Trabajo publicada en el Diario Oficial de la Federación el día 30 de noviembre del 2012.") 
      font = run3.font
      paragraphM.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraphN = doc.add_paragraph()
      run3 = paragraphN.add_run("El (la) Profesor (a) está de acuerdo en que el patrón le efectúe los descuentos de cuotas al Seguro Social que le correspondan.") 
      font = run3.font
      paragraphN.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraphO = doc.add_paragraph()
      run3 = paragraphO.add_run("DECIMA PRIMERA.- ")
      run3.bold = TRUE
      font = run3.font
      paragraphO.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run3 = paragraphO.add_run("El (a) Profesor (a) asistirá según el horario establecido en la cláusula séptima  conviniéndose que si trabaja el domingo  tendrá derecho a que se le pague una prima de un 25 % sobre su salario tabulado, quedando obligado a asistir a cursos de capacitación y desarrollo los días establecidos por la dirección técnica, la remuneración por estos últimos conceptos está incluida e integrado en la Cláusula Decima del presente Contrato.") 
      font = run3.font
      paragraphO.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraphP = doc.add_paragraph()
      run3 = paragraphP.add_run("DECIMA SEGUNDA.- ")
      run3.bold = TRUE
      font = run3.font
      paragraphP.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run3 = paragraphP.add_run("Son días de descanso obligatorio de acuerdo con el Articulo 74 de la Ley Federal del Trabajo, el 1° de Enero, 5 de Febrero, 21 de Marzo, 1° de Mayo, 16 de Septiembre, 20 de Noviembre, 25 de Diciembre y 1º de Diciembre de cada seis años, cuando corresponda a la transmisión del Poder Ejecutivo Federal, quedando prohibido que el (a) Profesor (a)  labore esos días, salvo permiso previo y por escrito del Patrón. El pago de estos días queda cubierto en la cantidad convenida como salario que aparece en la Cláusula Décima de este contrato. ") 
      font = run3.font
      paragraphP.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)
      
      paragraphQ = doc.add_paragraph()
      run3 = paragraphQ.add_run("Las partes aceptan los cambios a estas fechas para aprovechar los “Puentes” que determine la Ley Laboral.") 
      font = run3.font
      paragraphQ.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraphR = doc.add_paragraph()
      run3 = paragraphR.add_run("DECIMA TERCERA.- ")
      run3.bold = TRUE
      font = run3.font
      paragraphR.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run3 = paragraphR.add_run("El (a) Profesor (a) se compromete a sujetarse a los cursos de capacitación y adiestramiento a que se refieren los Artículos 153 A al 153 X de la Ley Federal del Trabajo.") 
      font = run3.font
      paragraphR.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraphS = doc.add_paragraph()
      run3 = paragraphS.add_run("DECIMA CUARTA.- ")
      run3.bold = TRUE
      font = run3.font
      paragraphS.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run3 = paragraphS.add_run("El (a) Profesor (a) gozará EXCLUSIVAMENTE de las VACACIONES Y DE LA PRIMA VACACIONAL que le correspondan con base en los Artículos 76 y 80 de la Ley Federal del Trabajo de acuerdo con su antigüedad en la escuela. Estas vacaciones serán disfrutadas exclusivamente durante los periodos de la última quincena de Diciembre de cada año y el periodo de Semana Santa tal como lo señala y ordena el Calendario Oficial de la SEP vigente.") 
      font = run3.font
      paragraphS.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraphT = doc.add_paragraph()
      run3 = paragraphT.add_run("El Profesor (a) se da por enterado (a) y de acuerdo en que los periodos de Julio y Agosto corresponden a RECESO DE CLASES para los alumnos, según dictamen de la Secretaría de Educación Pública y/o UNAM, y que por lo tanto, estos periodos en ningún caso corresponde a vacaciones para el personal docente y del Colegio en general.") 
      font = run3.font
      paragraphT.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraphU = doc.add_paragraph()
      run3 = paragraphU.add_run("DECIMA QUINTA.- ")
      run3.bold = TRUE
      font = run3.font
      paragraphU.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run3 = paragraphU.add_run("Todos los estudios, planes, programas, datos y en general cualquier documentación e información que el Profesor reciba o que elabore en el desempeño de sus servicios o por el encargo específico de la Escuela, serán propiedad de ésta última, en cuya virtud se obliga a devolverlos en el momento de ser requeridos para ello o al terminar este Contrato, o sea el " + final) 
      font = run3.font
      paragraphU.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraphV = doc.add_paragraph()
      run3 = paragraphV.add_run("DECIMA SEXTA.- ")
      run3.bold = TRUE
      font = run3.font
      paragraphV.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run3 = paragraphV.add_run("El (a) Profesor (a) está de acuerdo en no divulgar con ninguna persona o en otra Institución los datos, documentos, conocimientos e informes que haya obtenido con motivo de la prestación de sus servicios en la Escuela, ya que tienen el carácter de confidenciales, en caso de hacerlo será considerada su conducta como falta de probidad, además de las consecuencias legales que de esto origine.") 
      font = run3.font
      paragraphV.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraphW = doc.add_paragraph()
      run3 = paragraphW.add_run("DECIMA SEPTIMA.- ")
      run3.bold = TRUE
      font = run3.font
      paragraphW.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run3 = paragraphW.add_run("Los contratantes declaran que conocen el Reglamento Interior de Trabajo del Colegio, al cual se sujetarán en todas sus Cláusulas, por haberlo firmado protestando su estricto y legal cumplimiento.") 
      font = run3.font
      paragraphW.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraphX = doc.add_paragraph()
      run3 = paragraphX.add_run("DECIMA OCTAVA.- ")
      run3.bold = TRUE
      font = run3.font
      paragraphX.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run3 = paragraphX.add_run("El patrón reconoce al trabajador una antigüedad a partir del " + antiguedad + ". En todo lo no previsto, en el presente Contrato, se estará a las disposiciones de la Ley Federal del Trabajo vigente.") 
      font = run3.font
      paragraphX.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraphY = doc.add_paragraph()
      run3 = paragraphY.add_run("DECIMA NOVENA.- ")
      run3.bold = TRUE
      font = run3.font
      paragraphY.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run3 = paragraphY.add_run("El Instituto Francisco Possenti, A. C., le informa que sus datos personales, incluyendo los sensibles se utilizarán para identificar, informar, operar, gestionar y demás acciones que sean necesarias para la prestación de servicios laborales subordinados en el Instituto. El derecho de acceso, rectificación, cancelación oposición, limitación o la revocación de uso de sus datos personales, que para tal fin nos haya otorgado, a través de los procedimientos que hemos implementado, podrá solicitarse por escrito en dependencias gubernamentales como, la SEP, el SAT, IMSS, INFONAVIT, Secretaría del Trabajo etc., le informamos que si usted no manifiesta su oposición para que sus datos personales sean utilizados por el Instituto, significa que ha leído, entendido y aceptado los términos antes expuestos otorgando su consentimiento para ello.") 
      font = run3.font
      paragraphY.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraphZ = doc.add_paragraph()
      run3 = paragraphZ.add_run("VIGÉCIMA.- ")
      run3.bold = TRUE
      font = run3.font
      paragraphZ.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run3 = paragraphZ.add_run("Ambas partes acuerdan que en caso de suspensión de las actividades escolares por causas de fuerza mayor o caso fortuito se estará a las disposiciones que señalen las autoridades competentes.") 
      font = run3.font
      paragraphZ.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraphA1 = doc.add_paragraph()
      run3 = paragraphA1.add_run("VIGÉCIMA PRIMERA.- ")
      run3.bold = TRUE
      font = run3.font
      paragraphA1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run3 = paragraphA1.add_run("Ambas partes manifiestan que tienen pleno conocimiento  y cumplimiento de los alcances de la NOM-035-STPS-2018, factores de riesgo psicosocial en el  trabajo, identificación, análisis y prevención, publicada en el Diario  Oficial de la Federación el 23 de octubre del 2018, en términos de los dispuesto por los artículos 40, fracciones I y XI, de la Ley Orgánica  de la Administración Pública Federal; 512, 523, fracción  I, 524 y 527, último párrafo,  de  la Ley Federal  del Trabajo; 1º.,3º., fracción  XI, 38, fracción  II, 40, fracción VII, 41,47,  fracción IV,51, primer párrafo, 62,68, y 87 de la Ley Federal sobre  Metrología   y Normalización ;  28 del Reglamento de la  ley federal sobre Metrología  y Normalización; 5º., fracción III,7, fracciones I,II, III,IV,V, VII,IX,XI y XII, 8, fracciones I, III, V, VIII, X y XI,10, 32, fracción XI, 43,44, fracción VIII, y 55, del Reglamento  Federal de Seguridad y Salud en el Trabajo, y 5, fracción III, y 24 del Reglamento Interior de la Secretaría  del Trabajo y Previsión Social.") 
      font = run3.font
      paragraphA1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraphA2 = doc.add_paragraph()
      run3 = paragraphA2.add_run("VIGÉCIMA SEGUNDA.- ")
      run3.bold = TRUE
      font = run3.font
      paragraphA2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      run3 = paragraphA2.add_run("Ambas partes están de acuerdo que en caso de que hubiera una contingencia sanitaria conforme a los artículos 42 Bis, 429 fracción l y IV de la Ley Federal del Trabajo o declaración de emergencia sanitaria derivada de cualquier tipo de virus o bacteria, o situación derivada de cualquier tipo de “caso fortuito” o de “fuerza mayor” y como consecuencia de ello se tuviera que suspender la relación de trabajo o cerrar por tiempo determinado o indeterminado la fuente de trabajo, incluyendo la suspensión de la actividad escolar por tiempo definido o indefinido, Patrón y profesor(a) acuerdan que durante el periodo de dicha suspensión el salario quincenal integrado se ajustará de conformidad con las posibilidades económicas del Patrón, lo anterior con el objetivo de que los recursos económicos  (ingresos) de dicho patrón pueden ser repartidos de forma equitativa entre todos los trabajadores, incluyendo el hacer frente al pago de las obligaciones fiscales, gastos fijos y administrativos de la escuela y de seguridad social de cada Trabajador. A su vez y derivado de caso fortuito o fuerza mayor cuando las autoridades competentes nos obliguen a modificar en todo  o en parte el presente contrato se tomará en cuenta lo establecido en el decreto por el que se reforma el artículo 311 y se adiciona el capitulo XII Bis de la Ley Federal de Trabajo en materia de Teletrabajo publicado en el Diario Oficial de la Federación el día 11 de enero 2021 para aquellos casos o situaciones que apliquen en el presente contrato y relación de trabajo.") 
      font = run3.font
      paragraphA2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraphA3 = doc.add_paragraph()
      run3 = paragraphA3.add_run("Leído que fue por ambas partes este documento y sabedoras de las obligaciones que contraen, lo firman de conformidad por duplicado en el lugar y fecha señalados a continuación.") 
      font = run3.font
      paragraphA3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
      font.name = 'Arial'
      font.size = Pt(11)

      paragraphA2 = doc.add_paragraph()
      run3 = paragraphA2.add_run("CIUDAD DE MEXICO, A " + start)
      run3.bold = TRUE
      font = run3.font
      paragraphA2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
      font.name = 'Arial'
      font.size = Pt(11)

      table = doc.add_table(rows = 3, cols = 3)

      cell1 = table.cell(2,0)
      cell1.text = '      EL PATRÓN                        C. J. ANTONIO BARRIENTOS R. Representante Legal'
      run = cell1.paragraphs[0].runs[0]
      run.font.bold = True
      run.font.name = "Arial"
      run.font.size = Pt(9)
      cell1.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

      tc = cell1._tc
      tcPr = tc.get_or_add_tcPr()
      tcBorders = OxmlElement('w:tcBorders')
      top = OxmlElement('w:top')
      top.set(qn('w:val'), 'single')

      tcBorders.append(top)
      tcPr.append(tcBorders)
      
      cell3 = table.cell(2,2)
      cell3.text = 'EL (A) PROFESOR (A) ' + nombre
      run = cell3.paragraphs[0].runs[0]
      run.font.bold = True
      run.font.name = "Arial"
      run.font.size = Pt(9)
      cell3.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
      
      tc = cell3._tc
      tcPr = tc.get_or_add_tcPr()
      tcBorders = OxmlElement('w:tcBorders')
      top = OxmlElement('w:top')
      top.set(qn('w:val'), 'single')


      tcBorders.append(top)
      tcPr.append(tcBorders)

      doc.save("Contrato " + nombre + ".docx")
      

#Creamos interfas
root = Tk()
ex = StringVar()
hoja = StringVar()
root.title('Generación de contratos')

Label(root, text="Contratos", fg="darkblue", font=("Arial", 28, "bold")).pack()

#Pedimos datos necesarios de Excel
Label(root, text="Nombre del Excel",fg="black",font=("Arial", 16, "bold")).pack()
Entry(root, justify="center", textvariable=ex).pack()

Label(root, text="Nombre de la hoja", fg="black", font=("Arial", 16, "bold")).pack()
Entry(root, justify="center", textvariable=hoja).pack()

Button(root, text="Aceptar", command=crear).pack()

root.mainloop()